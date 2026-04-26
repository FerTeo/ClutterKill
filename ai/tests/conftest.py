"""
Shared pytest fixtures for the ingestion pipeline test suite.

Provides temporary test files (PDF, DOCX, PNG, TXT) and pre-configured
module instances so individual test files stay DRY.
"""

import asyncio
import os
import tempfile
from pathlib import Path
from unittest.mock import AsyncMock, MagicMock

import pytest
import pytest_asyncio

from PIL import Image, ImageDraw, ImageFont


# ── Async event loop ─────────────────────────────────────────────────────────

@pytest.fixture(scope="session")
def event_loop():
    """Create a single event loop for the whole test session."""
    loop = asyncio.new_event_loop()
    yield loop
    loop.close()


# ── Temporary test files ─────────────────────────────────────────────────────

@pytest.fixture
def tmp_dir():
    """Provide a temporary directory that is cleaned up after the test."""
    with tempfile.TemporaryDirectory() as d:
        yield Path(d)


@pytest.fixture
def sample_text_image(tmp_dir: Path) -> Path:
    """Create a PNG image containing readable text for OCR testing.

    Generates a 400x100 white image with black text 'Hello World 2024'
    at 72 DPI — a baseline for verifying OCR accuracy.
    """
    img = Image.new("RGB", (400, 100), color="white")
    draw = ImageDraw.Draw(img)

    # Use a basic font — available on all systems
    try:
        font = ImageFont.truetype("arial.ttf", 36)
    except (OSError, IOError):
        font = ImageFont.load_default()

    draw.text((20, 30), "Hello World 2024", fill="black", font=font)

    path = tmp_dir / "test_image.png"
    img.save(str(path), dpi=(72, 72))
    return path


@pytest.fixture
def sample_docx(tmp_dir: Path) -> Path:
    """Create a minimal .docx file with known content for parser testing."""
    from docx import Document

    doc = Document()
    doc.core_properties.author = "Test Author"
    doc.core_properties.title = "Test Title"

    doc.add_paragraph("This is the first paragraph of the test document.")
    doc.add_paragraph("Second paragraph with important information.")
    doc.add_paragraph("Third paragraph: final conclusion and summary.")

    # Add a small table
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Score"
    table.cell(1, 1).text = "95"

    path = tmp_dir / "test_document.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def sample_empty_image(tmp_dir: Path) -> Path:
    """Create a blank white image (no text) for edge-case testing."""
    img = Image.new("RGB", (200, 200), color="white")
    path = tmp_dir / "blank.png"
    img.save(str(path))
    return path


@pytest.fixture
def mock_ocr_engine():
    """A mock OCR engine that returns predictable text."""
    engine = MagicMock()
    engine.extract_text = AsyncMock(return_value="Mocked OCR text output")
    engine.extract_text_from_path = AsyncMock(return_value="Mocked OCR text output")
    return engine

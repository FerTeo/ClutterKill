"""
DOCX Document Parser

Extracts text from Microsoft Word (``.docx``) files using the
``python-docx`` library.

Capabilities
------------
* **Paragraph extraction** — each paragraph becomes a logical block.
* **Table extraction** — tables are rendered as tab-separated text so
  the LLM can still reason about tabular data.
* **Rich metadata** — author, title, creation date, last modified, and
  word count are pulled from the OOXML core properties.

Design notes
------------
* DOCX files don't have "pages" in the same sense as PDFs (page breaks
  depend on the rendering engine), so ``page_count`` is estimated from
  the ``python-docx`` built-in properties when available, and
  ``raw_pages`` contains a single entry with the full text.
* The heavy I/O is wrapped in :func:`asyncio.to_thread`.
"""

from __future__ import annotations

import asyncio
import logging
from pathlib import Path

from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError

from agents.parsers.base import BaseDocumentParser, ParsedDocument

logger = logging.getLogger(__name__)


# ── Helpers ──────────────────────────────────────────────────────────────────

def _extract_docx_metadata(doc: DocxDocument) -> dict:
    """Extract core properties from the OOXML package.

    These map directly to the Dublin Core / OOXML core-properties XML
    inside the ``.docx`` ZIP archive.
    """
    props = doc.core_properties
    return {
        "author": props.author,
        "title": props.title,
        "subject": props.subject,
        "category": props.category,
        "keywords": props.keywords,
        "created": str(props.created) if props.created else None,
        "modified": str(props.modified) if props.modified else None,
        "last_modified_by": props.last_modified_by,
        "revision": props.revision,
    }


def _extract_tables(doc: DocxDocument) -> list[str]:
    """Render every table in the document as tab-separated text.

    Each table becomes a single string where rows are separated by
    newlines and cells by tabs — a format LLMs handle well.
    """
    table_texts: list[str] = []
    for table in doc.tables:
        rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append("\t".join(cells))
        table_texts.append("\n".join(rows))
    return table_texts


def _sync_extract(file_path: Path) -> tuple[str, dict, int]:
    """Synchronous extraction — runs inside a worker thread.

    Returns
    -------
    (full_text, metadata, paragraph_count)
    """
    try:
        doc = DocxDocument(str(file_path))
    except PackageNotFoundError as exc:
        raise ValueError(
            f"File '{file_path.name}' is not a valid DOCX package."
        ) from exc

    # Paragraphs
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    # Tables (appended after paragraphs)
    tables = _extract_tables(doc)

    # Combine
    parts: list[str] = []
    if paragraphs:
        parts.append("\n\n".join(paragraphs))
    if tables:
        parts.append("\n\n--- TABLE ---\n" + "\n\n--- TABLE ---\n".join(tables))

    full_text = "\n\n".join(parts)

    metadata = _extract_docx_metadata(doc)

    return full_text, metadata, len(paragraphs)


# ── Parser ───────────────────────────────────────────────────────────────────

class DocxParser(BaseDocumentParser):
    """Parse ``.docx`` files using ``python-docx``.

    Extracts paragraph text, table content, and rich metadata from
    the OOXML core properties.

    The ``max_pages`` parameter is accepted for interface consistency
    but has limited effect: DOCX files do not have a reliable page
    concept outside of a rendering engine.  If ``max_pages`` is set,
    the parser will truncate after approximately
    ``max_pages × 40`` paragraphs (rough heuristic).

    Examples
    --------
    >>> parser = DocxParser()
    >>> doc = await parser.parse(Path("contract.docx"))
    >>> print(doc.metadata.get("author"))
    'John Doe'
    """

    # Rough estimate: ~40 paragraphs per page in a typical document.
    _PARAGRAPHS_PER_PAGE: int = 40

    async def parse(
        self,
        file_path: Path,
        max_pages: int | None = None,
    ) -> ParsedDocument:
        """Extract text and metadata from a DOCX file.

        Parameters
        ----------
        file_path : Path
            Absolute path to the ``.docx`` file.
        max_pages : int or None
            Rough page cap (heuristic: 40 paragraphs ≈ 1 page).
            ``None`` reads the entire document.

        Returns
        -------
        ParsedDocument
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"DOCX file not found: {file_path}")

        logger.info("Parsing DOCX '%s'", file_path.name)

        full_text, metadata, para_count = await asyncio.to_thread(
            _sync_extract, file_path
        )

        # Apply rough page truncation if requested
        if max_pages is not None:
            max_chars = max_pages * self._PARAGRAPHS_PER_PAGE * 80  # ~80 chars/para
            if len(full_text) > max_chars:
                full_text = full_text[:max_chars]
                logger.info(
                    "DOCX '%s' truncated to ~%d pages (%d chars)",
                    file_path.name,
                    max_pages,
                    max_chars,
                )

        # Estimate page count from paragraph count
        estimated_pages = max(1, para_count // self._PARAGRAPHS_PER_PAGE)

        logger.info(
            "DOCX '%s': ~%d paragraphs, %d chars extracted, author='%s'",
            file_path.name,
            para_count,
            len(full_text),
            metadata.get("author", "unknown"),
        )

        return ParsedDocument(
            text=full_text,
            source_path=str(file_path.resolve()),
            mime_type=(
                "application/vnd.openxmlformats-officedocument"
                ".wordprocessingml.document"
            ),
            page_count=estimated_pages,
            is_ocr_result=False,
            metadata=metadata,
            raw_pages=[full_text],  # Single "page" for DOCX
        )

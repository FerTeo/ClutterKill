"""
generate_mock_data.py — scripts/generate_mock_data.py
# mypy: ignore-errors

Generează 20 de fișiere neorganizate în test_data/source/:
  - 7 fișiere PDF (facturi / documente de curs) folosind fpdf
  - 7 imagini PNG (scanări simulate) folosind Pillow
  - 6 documente Word (.docx) folosind python-docx

Toate fișierele conțin cuvinte cheie: „Factură", „Curs", „Semestru".
"""

import random
import sys
from pathlib import Path

# ---------------------------------------------------------------------------
# Asigurăm că rădăcina proiectului este în sys.path (rulare din orice dir)
# ---------------------------------------------------------------------------
PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

from fpdf import FPDF  # noqa: E402
from PIL import Image, ImageDraw, ImageFont  # noqa: E402
from docx import Document  # noqa: E402
from docx.shared import Pt  # noqa: E402

# ---------------------------------------------------------------------------
# Directorul destinație
# ---------------------------------------------------------------------------
OUTPUT_DIR = PROJECT_ROOT / "test_data" / "source"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# ---------------------------------------------------------------------------
# Date de test (variabile pentru a genera fișiere distincte)
# ---------------------------------------------------------------------------
FURNIZORI = [
    "SC TechSoft SRL",
    "Universitatea Politehnica",
    "EduConsult SA",
    "InfoSystems SRL",
    "AcademiaPro",
    "DataVision SRL",
    "SmartLearn SRL",
]

CURSURI = [
    "Metode de Dezvoltare a Softului",
    "Structuri de Date si Algoritmi",
    "Baze de Date Avansate",
    "Ingineria Sistemelor Software",
    "Retele de Calculatoare",
    "Inteligenta Artificiala",
    "Programare Orientata Obiect",
]

SEMESTRE = ["Semestrul 1", "Semestrul 2", "Semestrul I", "Semestrul II"]

DATE = [
    "10.01.2025",
    "15.02.2025",
    "03.03.2025",
    "22.04.2025",
    "07.05.2025",
    "18.06.2025",
    "30.09.2025",
    "11.10.2025",
    "25.11.2025",
    "02.12.2025",
    "14.01.2026",
    "28.02.2026",
]

PRODUSE = [
    ("Licenta software educational", 1, 450.00),
    ("Suport curs semestrial", 2, 200.00),
    ("Acces platforma e-learning", 1, 120.00),
    ("Manual digital Semestrul 2", 3, 75.00),
    ("Consultanta implementare proiect", 5, 150.00),
    ("Abonament resurse academice", 1, 350.00),
    ("Kit laborator programare", 2, 180.00),
]


def _rand(lst):
    return random.choice(lst)


def _p(text: str) -> str:
    """Strip orice caracter non-latin-1 (necesar pentru fpdf v1)."""
    return text.encode("latin-1", errors="ignore").decode("latin-1")


# ===========================================================================
# 1. PDF — 7 fișiere
# ===========================================================================


def _make_pdf(index: int) -> Path:
    """Generează o factură PDF cu fpdf."""
    furnizor = FURNIZORI[index % len(FURNIZORI)]
    curs = CURSURI[index % len(CURSURI)]
    semestru = _rand(SEMESTRE)
    data = DATE[index % len(DATE)]
    nr_factura = f"FCT-{2025 + index % 2}-{1000 + index * 37}"
    produs_name, qty, pret_unit = PRODUSE[index % len(PRODUSE)]
    total = qty * pret_unit
    tva = round(total * 0.19, 2)

    pdf = FPDF()
    pdf.add_page()

    # Titlu
    pdf.set_font("Helvetica", "B", 18)
    pdf.set_fill_color(30, 60, 120)
    pdf.set_text_color(255, 255, 255)
    pdf.cell(0, 12, _p("FACTURA FISCALA"), ln=True, align="C", fill=True)
    pdf.ln(4)

    # Antet document
    pdf.set_text_color(0, 0, 0)
    pdf.set_font("Helvetica", "", 11)
    pdf.cell(90, 8, _p(f"Numar Factura: {nr_factura}"), ln=False)
    pdf.cell(0, 8, _p(f"Data: {data}"), ln=True)
    pdf.cell(90, 8, _p(f"Furnizor: {furnizor}"), ln=False)
    pdf.cell(0, 8, _p("Client: Universitatea Tehnica"), ln=True)
    pdf.ln(4)

    # Linie separatoare
    pdf.set_draw_color(100, 100, 100)
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(4)

    # Detalii curs / semestru
    pdf.set_font("Helvetica", "B", 11)
    pdf.cell(0, 8, _p(f"Curs: {curs}"), ln=True)
    pdf.cell(0, 8, _p(f"Perioada: {semestru} 2024-2025"), ln=True)
    pdf.ln(4)

    # Tabel produse
    pdf.set_fill_color(200, 215, 240)
    pdf.set_font("Helvetica", "B", 10)
    pdf.cell(80, 8, _p("Descriere produs/serviciu"), border=1, fill=True)
    pdf.cell(20, 8, _p("Cant."), border=1, fill=True, align="C")
    pdf.cell(35, 8, _p("Pret unit. (RON)"), border=1, fill=True, align="C")
    pdf.cell(35, 8, _p("Total (RON)"), border=1, fill=True, align="C", ln=True)

    pdf.set_font("Helvetica", "", 10)
    pdf.cell(80, 8, _p(produs_name), border=1)
    pdf.cell(20, 8, str(qty), border=1, align="C")
    pdf.cell(35, 8, f"{pret_unit:.2f}", border=1, align="R")
    pdf.cell(35, 8, f"{total:.2f}", border=1, align="R", ln=True)

    pdf.ln(2)
    pdf.set_font("Helvetica", "", 10)
    pdf.cell(135, 8, _p("Subtotal:"), align="R")
    pdf.cell(35, 8, f"{total:.2f} RON", align="R", ln=True)
    pdf.cell(135, 8, _p("TVA (19%):"), align="R")
    pdf.cell(35, 8, f"{tva:.2f} RON", align="R", ln=True)
    pdf.set_font("Helvetica", "B", 11)
    pdf.cell(135, 8, _p("TOTAL DE PLATA:"), align="R")
    pdf.cell(35, 8, f"{total + tva:.2f} RON", align="R", ln=True)

    pdf.ln(8)
    pdf.set_font("Helvetica", "I", 9)
    pdf.set_text_color(80, 80, 80)
    pdf.cell(
        0,
        6,
        _p(f"Aceasta Factura este generata automat. {semestru} - {curs}."),
        ln=True,
        align="C",
    )

    filename = OUTPUT_DIR / f"factura_{nr_factura.replace('-', '_')}.pdf"
    pdf.output(str(filename))
    return filename


# ===========================================================================
# 2. Imagini PNG — 7 fișiere (scanări simulate)
# ===========================================================================


def _make_image(index: int) -> Path:
    """Generează o imagine PNG simulând un document scanat."""
    furnizor = FURNIZORI[index % len(FURNIZORI)]
    curs = CURSURI[index % len(CURSURI)]
    semestru = SEMESTRE[index % len(SEMESTRE)]
    data = DATE[(index + 3) % len(DATE)]
    nr = f"IMG-{1000 + index * 53}"

    # Fundal ușor gălbui (hârtie veche)
    bg_color = (252, 248, 230) if index % 2 == 0 else (255, 255, 255)
    img = Image.new("RGB", (800, 600), color=bg_color)
    draw = ImageDraw.Draw(img)

    # Încercăm un font mai mare; fallback la default
    try:
        font_title = ImageFont.truetype("arial.ttf", 26)
        font_body = ImageFont.truetype("arial.ttf", 18)
        font_small = ImageFont.truetype("arial.ttf", 14)
    except OSError:
        font_title = ImageFont.load_default()
        font_body = font_title
        font_small = font_title

    text_color = (20, 20, 20)
    accent = (30, 60, 120)

    # Header
    draw.rectangle([0, 0, 800, 55], fill=accent)
    draw.text(
        (20, 12),
        "FACTURA FISCALA / DOCUMENT CURS",
        font=font_title,
        fill=(255, 255, 255),
    )

    y = 75
    draw.text((20, y), f"Nr. Document: {nr}", font=font_body, fill=text_color)
    y += 30
    draw.text((20, y), f"Data emiterii: {data}", font=font_body, fill=text_color)
    y += 30
    draw.text((20, y), f"Furnizor: {furnizor}", font=font_body, fill=text_color)
    y += 30
    draw.text((20, y), f"Curs: {curs}", font=font_body, fill=accent)
    y += 30
    draw.text(
        (20, y), f"Semestru: {semestru} 2024-2025", font=font_body, fill=text_color
    )
    y += 35

    # Linie separator
    draw.line([(20, y), (780, y)], fill=(150, 150, 150), width=2)
    y += 15

    draw.text((20, y), "Descriere:", font=font_body, fill=text_color)
    y += 28
    draw.text(
        (30, y),
        "Acest document atesta plata serviciilor educationale aferente cursului",
        font=font_small,
        fill=text_color,
    )
    y += 22
    draw.text(
        (30, y),
        f'"{curs}" desfasurat in {semestru}.',
        font=font_small,
        fill=text_color,
    )
    y += 30

    total = round(random.uniform(100, 800), 2)
    draw.text(
        (20, y),
        f"Suma totala de plata: {total} RON (incl. TVA 19%)",
        font=font_body,
        fill=text_color,
    )
    y += 35

    draw.line([(20, y), (780, y)], fill=(150, 150, 150), width=1)
    y += 10
    draw.text(
        (20, y),
        "Semnat si stampilat. Document valid in conformitate cu legislatia in vigoare.",
        font=font_small,
        fill=(100, 100, 100),
    )

    # Adăugăm un zgomot ușor pentru a simula scanarea
    if index % 3 == 0:
        pixels = img.load()
        for _ in range(500):
            rx = random.randint(0, 799)
            ry = random.randint(56, 599)
            pixels[rx, ry] = (random.randint(180, 220),) * 3

    filename = OUTPUT_DIR / f"scan_document_{nr}.png"
    img.save(str(filename))
    return filename


# ===========================================================================
# 3. Word Documents (.docx) — 6 fișiere
# ===========================================================================


def _make_docx(index: int) -> Path:
    """Generează un document Word cu suport de curs sau factură consultanță."""
    furnizor = FURNIZORI[index % len(FURNIZORI)]
    curs = CURSURI[index % len(CURSURI)]
    semestru = SEMESTRE[index % len(SEMESTRE)]
    data = DATE[(index + 5) % len(DATE)]
    nr = f"DOC-{2025 + index % 2}-{500 + index * 61}"

    doc = Document()

    # Stil titlu
    title = doc.add_heading(level=1)
    title.clear()
    run = title.add_run(f"Factură / Suport Curs — {semestru}")
    run.font.size = Pt(18)

    doc.add_paragraph(f"Număr document: {nr}")
    doc.add_paragraph(f"Data emiterii: {data}")
    doc.add_paragraph(f"Furnizor: {furnizor}")
    doc.add_paragraph("")

    doc.add_heading("Detalii Curs", level=2)
    doc.add_paragraph(f"Curs: {curs}")
    doc.add_paragraph(f"Semestru: {semestru} 2024-2025")
    doc.add_paragraph(
        f"Acest document reprezintă materialul suport pentru cursul «{curs}» "
        f"organizat în cadrul {semestru} al anului universitar 2024-2025."
    )

    doc.add_heading("Factură Consultanță", level=2)
    descriere_para = doc.add_paragraph()
    descriere_para.add_run(
        f"Factura nr. {nr} emisă de {furnizor} pentru servicii de consultanță "
        f"și suport educațional aferent cursului {curs}, {semestru}."
    )

    # Tabel produse
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Descriere"
    hdr[1].text = "Cantitate"
    hdr[2].text = "Preț unitar (RON)"
    hdr[3].text = "Total (RON)"

    produs_name, qty, pret_unit = PRODUSE[index % len(PRODUSE)]
    total = qty * pret_unit
    row = table.add_row().cells
    row[0].text = produs_name
    row[1].text = str(qty)
    row[2].text = f"{pret_unit:.2f}"
    row[3].text = f"{total:.2f}"

    doc.add_paragraph("")
    total_tva = round(total * 1.19, 2)
    doc.add_paragraph(f"Total de plată (incl. TVA 19%): {total_tva:.2f} RON")

    doc.add_heading("Semnătură și Confirmare", level=2)
    doc.add_paragraph(
        "Document semnat electronic. Valid pentru Semestrul în curs. "
        f"Curs desfășurat conform planului universitar — {semestru}."
    )

    filename = OUTPUT_DIR / f"document_{nr.replace('-', '_')}.docx"
    doc.save(str(filename))
    return filename


# ===========================================================================
# Main
# ===========================================================================


def main():
    print(f"[INFO] Director destinatie: {OUTPUT_DIR}")
    print("[INFO] Generare fisiere mock...\n")

    generated = []

    # 7 PDF-uri
    for i in range(7):
        path = _make_pdf(i)
        generated.append(path)
        print(f"  [PDF {i + 1}/7]  {path.name}")

    # 7 Imagini PNG
    for i in range(7):
        path = _make_image(i)
        generated.append(path)
        print(f"  [IMG {i + 1}/7]  {path.name}")

    # 6 Documente Word
    for i in range(6):
        path = _make_docx(i)
        generated.append(path)
        print(f"  [DOC {i + 1}/6]  {path.name}")

    print(f"\n[DONE] Generat {len(generated)} fisiere in {OUTPUT_DIR}")
    return generated


if __name__ == "__main__":
    main()
